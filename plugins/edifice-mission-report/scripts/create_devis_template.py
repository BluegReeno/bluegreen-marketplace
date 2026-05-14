#!/usr/bin/env python3
"""
create_devis_template.py

Creates templates/ic-ingenieurs/devis.docx from scratch with IC Ingénieurs branding.

Run with:
    uv run --with "python-docx>=1.1" python3 scripts/create_devis_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

PLUGIN_DIR = Path(__file__).resolve().parent.parent
OUTPUT = PLUGIN_DIR / "templates" / "ic-ingenieurs" / "devis.docx"

IC_BLUE_HEX = "1F3A5F"
IC_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
IC_ORANGE_HEX = "E86A1A"
IC_ORANGE = RGBColor(0xE8, 0x6A, 0x1A)
IC_ROW_ALT_HEX = "F0F4F8"
WHITE_HEX = "FFFFFF"
LIGHT_GREY_HEX = "F5F5F5"


def _set_cell_shading(cell, hex_color: str) -> None:
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shd)


def _set_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _set_col_widths(table, widths_cm: list) -> None:
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for col in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(col)
    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w * 567)))
        tblGrid.append(gridCol)


def _add_run(para, text: str, bold=False, italic=False, color=None, size_pt=None):
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def add_section_heading(doc: Document, title: str, level: int = 1) -> None:
    p = doc.add_heading(title, level=level)
    for run in p.runs:
        run.font.color.rgb = IC_BLUE


def add_2col_table(doc: Document, rows_data: list, widths_cm=(5.0, 11.0)) -> None:
    """Add a 2-col key-value table. rows_data = list of (label, value_tag)."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    _set_col_widths(table, list(widths_cm))
    for i, (label, tag) in enumerate(rows_data):
        row = table.rows[i]
        p_label = row.cells[0].paragraphs[0]
        _add_run(p_label, label, bold=True, size_pt=9)
        _set_cell_shading(row.cells[0], IC_ROW_ALT_HEX if i % 2 == 0 else WHITE_HEX)
        p_val = row.cells[1].paragraphs[0]
        p_val.add_run(tag)


def add_loop_table(doc: Document, header_cols: list, loop_var: str, fields: list,
                   widths_cm: list) -> None:
    """
    Add a table with a docxtpl {%tr for %} loop.
    header_cols: column header labels
    loop_var: e.g. "d in documents_fournis"
    fields: list of {{ d.field }} tags (one per col)
    """
    n_cols = len(header_cols)
    # Header row + loop start row + data row + loop end row
    table = doc.add_table(rows=4, cols=n_cols)
    table.style = "Table Grid"
    _set_col_widths(table, widths_cm)

    # Row 0: column headers (IC_BLUE bg, white text)
    r0 = table.rows[0]
    for j, h in enumerate(header_cols):
        cell = r0.cells[j]
        _set_cell_shading(cell, IC_BLUE_HEX)
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=9)

    # Row 1: loop start (minimal height)
    r1 = table.rows[1]
    _set_row_height(r1, 100)
    r1.cells[0].paragraphs[0].add_run(f"{{%tr for {loop_var} %}}")

    # Row 2: data row
    r2 = table.rows[2]
    for j, field in enumerate(fields):
        r2.cells[j].paragraphs[0].add_run(field)

    # Row 3: loop end (minimal height)
    r3 = table.rows[3]
    _set_row_height(r3, 100)
    r3.cells[0].paragraphs[0].add_run("{%tr endfor %}")


def add_cover(doc: Document) -> None:
    # IC_BLUE header bar
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [16.0])
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, IC_BLUE_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "IC INGÉNIEURS CONSEILS", bold=True,
             color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=14)

    doc.add_paragraph()

    # Document type badge
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_type, "RAPPORT PRÉLIMINAIRE — DEMANDE DE DEVIS",
             bold=True, color=IC_ORANGE, size_pt=13)

    doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, "{{ titre_service }}", bold=True, color=IC_BLUE, size_pt=16)

    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    _set_col_widths(meta_table, [5.0, 11.0])
    meta_rows = [
        ("Client\xa0:", "{{ client }}"),
        ("Adresse\xa0:", "{{ adresse }}"),
        ("Date visite\xa0:", "{{ date_visite }}"),
        ("Technicien\xa0:", "{{ technicien }}"),
    ]
    for i, (label, tag) in enumerate(meta_rows):
        row = meta_table.rows[i]
        p_l = row.cells[0].paragraphs[0]
        _add_run(p_l, label, bold=True, size_pt=10)
        _set_cell_shading(row.cells[0], IC_ROW_ALT_HEX)
        row.cells[1].paragraphs[0].add_run(tag)

    doc.add_page_break()


def main() -> None:
    print(f"Output: {OUTPUT}")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # --- Cover ---
    add_cover(doc)

    # --- Section 1: Contexte client ---
    add_section_heading(doc, "1. Contexte client")
    add_2col_table(doc, [
        ("Client\xa0:", "{{ client }}"),
        ("Type d'acteur\xa0:", "{{ type_acteur }}"),
        ("Interlocuteur\xa0:", "{{ interlocuteur_nom }}"),
        ("Rôle\xa0:", "{{ interlocuteur_role }}"),
        ("Contact\xa0:", "{{ interlocuteur_contact }}"),
    ])
    doc.add_paragraph()

    # --- Section 2: Mission ---
    add_section_heading(doc, "2. Mission")
    add_2col_table(doc, [
        ("Type de mission\xa0:", "{{ type_mission }}"),
        ("Déclencheur\xa0:", "{{ declencheur }}"),
        ("Livrable attendu\xa0:", "{{ livrable }}"),
        ("Urgence\xa0:", "{{ urgence }}"),
    ])
    doc.add_paragraph()

    # --- Section 3: Bâtiment ---
    add_section_heading(doc, "3. Bâtiment")
    add_2col_table(doc, [
        ("Adresse\xa0:", "{{ adresse }}"),
        ("Type de bâtiment\xa0:", "{{ type_batiment }}"),
        ("Année de construction\xa0:", "{{ annee_construction }}"),
        ("Nombre d'étages\xa0:", "{{ nb_etages }}"),
        ("Description\xa0:", "{{ description_batiment }}"),
    ])
    doc.add_paragraph()

    # --- Section 4: Documents fournis ---
    add_section_heading(doc, "4. Documents fournis")
    add_loop_table(
        doc,
        header_cols=["Document", "Fourni"],
        loop_var="d in documents_fournis",
        fields=["{{ d.document }}", "{{ d.fourni_str }}"],
        widths_cm=[12.0, 4.0],
    )
    doc.add_paragraph()

    # --- Section 5: Observations terrain ---
    add_section_heading(doc, "5. Observations terrain")
    add_loop_table(
        doc,
        header_cols=["Localisation", "Désordre", "Données clés", "Réf. photo"],
        loop_var="obs in observations",
        fields=["{{ obs.localisation }}", "{{ obs.desordre }}",
                "{{ obs.donnees_cles }}", "{{ obs.ref_photo }}"],
        widths_cm=[3.5, 5.5, 4.0, 3.0],
    )
    doc.add_paragraph()

    # --- Section 6: Proposition de mission ---
    add_section_heading(doc, "6. Proposition de mission")
    add_2col_table(doc, [
        ("Proposition\xa0:", "{{ proposition_mission }}"),
        ("Incertitudes / hypothèses\xa0:", "{{ incertitudes }}"),
    ], widths_cm=(5.0, 11.0))
    doc.add_paragraph()

    # --- Section 7: Chiffrage estimatif ---
    add_section_heading(doc, "7. Chiffrage estimatif")
    add_loop_table(
        doc,
        header_cols=["Prestation", "Nb heures", "Montant HT (€)"],
        loop_var="ligne in chiffrage",
        fields=["{{ ligne.prestation }}", "{{ ligne.nb_heures }}", "{{ ligne.montant_ht }}"],
        widths_cm=[9.0, 3.5, 3.5],
    )
    doc.add_paragraph()

    # --- Validation ---
    add_section_heading(doc, "Validation", level=2)
    val_table = doc.add_table(rows=2, cols=3)
    val_table.style = "Table Grid"
    _set_col_widths(val_table, [5.5, 5.5, 5.0])
    headers = ["Établi par", "Date de visite", "Date d'envoi"]
    for j, h in enumerate(headers):
        cell = val_table.rows[0].cells[j]
        _set_cell_shading(cell, IC_BLUE_HEX)
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=9)
    val_vals = ["{{ technicien }}", "{{ date_visite }}", "{{ date_envoi }}"]
    for j, v in enumerate(val_vals):
        val_table.rows[1].cells[j].paragraphs[0].add_run(v)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"✅ Saved: {OUTPUT}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
