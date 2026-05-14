#!/usr/bin/env python3
"""
update_suivi_chantier_branding.py

Loads templates/ic-ingenieurs/suivi_chantier.docx, replaces the plain cover page
with IC Ingénieurs branding, updates Heading 1/2 styles, saves in place.

Run with:
    uv run --with "python-docx>=1.1" python3 scripts/update_suivi_chantier_branding.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

PLUGIN_DIR = Path(__file__).resolve().parent.parent
TEMPLATE = PLUGIN_DIR / "templates" / "ic-ingenieurs" / "suivi_chantier.docx"

IC_BLUE_HEX = "1F3A5F"
IC_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
IC_ORANGE_HEX = "E86A1A"
IC_ORANGE = RGBColor(0xE8, 0x6A, 0x1A)
IC_ROW_ALT_HEX = "F0F4F8"
WHITE_HEX = "FFFFFF"


def _set_cell_shading(cell, hex_color: str) -> None:
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shd)


def _set_col_widths(table, widths_cm: list) -> None:
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for col in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(col)
    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w * 567)))
        tblGrid.append(gridCol)


def _add_run(para, text: str, bold=False, color=None, size_pt=None):
    run = para.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def build_cover_elements(doc: Document) -> list:
    """Build IC-branded cover elements; return list of XML elements to insert."""
    elements = []

    # --- IC_BLUE header bar table ---
    tbl_header = doc.add_table(rows=1, cols=1)
    tbl_header.style = "Table Grid"
    _set_col_widths(tbl_header, [16.0])
    cell = tbl_header.rows[0].cells[0]
    _set_cell_shading(cell, IC_BLUE_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "IC INGÉNIEURS CONSEILS", bold=True,
             color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=14)
    elements.append(tbl_header._tbl)

    # Blank paragraph
    blank = OxmlElement("w:p")
    elements.append(blank)

    # Document type
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_type, "COMPTE RENDU DE VISITE DE CHANTIER",
             bold=True, color=IC_ORANGE, size_pt=12)
    elements.append(p_type._element)

    # Service title
    blank2 = OxmlElement("w:p")
    elements.append(blank2)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_title, "{{ titre_service }}", bold=True, color=IC_BLUE, size_pt=16)
    elements.append(p_title._element)

    # Blank
    blank3 = OxmlElement("w:p")
    elements.append(blank3)

    # Meta table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Table Grid"
    _set_col_widths(meta_table, [5.0, 11.0])
    meta_rows = [
        ("Chantier\xa0:", "{{ residence }}"),
        ("Bâtiments visités\xa0:", "{{ batiments_visites }}"),
        ("Adresse\xa0:", "{{ adresse }}"),
        ("Réf. dossier\xa0:", "{{ ref_dossier }}"),
        ("Date de visite\xa0:", "{{ date_visite }}"),
    ]
    for i, (label, tag) in enumerate(meta_rows):
        row = meta_table.rows[i]
        p_l = row.cells[0].paragraphs[0]
        _add_run(p_l, label, bold=True, size_pt=10)
        _set_cell_shading(row.cells[0], IC_ROW_ALT_HEX if i % 2 == 0 else WHITE_HEX)
        row.cells[1].paragraphs[0].add_run(tag)
    elements.append(meta_table._tbl)

    # Blank
    blank4 = OxmlElement("w:p")
    elements.append(blank4)

    # Participants table (4 cols: Nom | Fonction | Entreprise | Contact)
    p_label = doc.add_paragraph()
    _add_run(p_label, "Participants :", bold=True, color=IC_BLUE, size_pt=10)
    elements.append(p_label._element)

    ptable = doc.add_table(rows=1, cols=4)
    ptable.style = "Table Grid"
    _set_col_widths(ptable, [4.0, 4.0, 4.5, 3.5])
    headers = ["Nom", "Fonction", "Entreprise", "Contact"]
    for j, h in enumerate(headers):
        cell = ptable.rows[0].cells[j]
        _set_cell_shading(cell, IC_BLUE_HEX)
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=9)
    elements.append(ptable._tbl)

    # Page break
    pb = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    pb_br = OxmlElement("w:br")
    pb_br.set(qn("w:type"), "page")
    pb_r.append(pb_br)
    pb.append(pb_r)
    elements.append(pb)

    return elements


def update_heading_styles(doc: Document) -> None:
    try:
        h1 = doc.styles["Heading 1"]
        h1.font.color.rgb = IC_BLUE
        h1.font.bold = True
        h1.font.size = Pt(13)
    except KeyError:
        pass
    try:
        h2 = doc.styles["Heading 2"]
        h2.font.color.rgb = IC_BLUE
        h2.font.bold = False
        h2.font.size = Pt(11)
    except KeyError:
        pass


def main() -> None:
    print(f"Template: {TEMPLATE}")

    doc = Document(str(TEMPLATE))
    body = doc.element.body
    children = list(body)

    # Find the first Heading 1 — everything before it is the old cover
    first_heading_el = None
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            first_heading_el = p._element
            break

    if first_heading_el is None:
        raise SystemExit("ERROR: No Heading paragraph found in template.")

    heading_idx = children.index(first_heading_el)
    print(f"Removing old cover: body elements [0:{heading_idx}]")

    # Remove old cover elements
    to_remove = children[:heading_idx]
    for el in to_remove:
        body.remove(el)

    # Rebuild children list
    children = list(body)
    heading_idx = children.index(first_heading_el)

    # Build new IC-branded cover
    new_elements = build_cover_elements(doc)
    # They were added to end of doc during build — remove them from there
    # then insert at correct position
    # (build_cover_elements returns elements that are already orphaned or moved)
    # Insert in reverse order before the heading
    for el in reversed(new_elements):
        # Remove from wherever it currently is (if attached)
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
        body.insert(heading_idx, el)

    # Update heading styles
    update_heading_styles(doc)

    doc.save(str(TEMPLATE))
    print(f"✅ Saved: {TEMPLATE}")

    # Verify Jinja2 tags preserved
    all_tags = []
    for p in doc.paragraphs:
        if "{{" in p.text or "{%" in p.text:
            all_tags.append(p.text[:60])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "{{" in cell.text or "{%" in cell.text:
                    all_tags.append(cell.text[:60])
    print(f"   Jinja2 tags preserved: {len(all_tags)}")
    for tag in all_tags[:10]:
        print(f"     {tag!r}")


if __name__ == "__main__":
    main()
