#!/usr/bin/env python3
"""
prepare_diagnostic_template.py

Loads rapport-diag-template.docx (IC branded, from MyClaudeSkills),
removes the static observations section, injects a docxtpl disorder loop table,
saves as templates/ic-ingenieurs/diagnostic.docx.

Run with:
    uv run --with "python-docx>=1.1" python3 scripts/prepare_diagnostic_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
PLUGIN_DIR = SCRIPT_DIR.parent

SOURCE = Path("/Users/renaud/Projects/MyClaudeSkills/document-generator/organizations/ic-ingenieurs/templates/rapport-diag-template.docx")
OUTPUT = PLUGIN_DIR / "templates" / "ic-ingenieurs" / "diagnostic.docx"

IC_BLUE_HEX = "1F3A5F"
IC_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
IC_ORANGE_HEX = "E86A1A"
WHITE_HEX = "FFFFFF"


def _set_cell_shading(cell, hex_color: str) -> None:
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shd)


def _set_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _add_run(para, text: str, bold=False, color=None, size_pt=None) -> None:
    run = para.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)


def _set_col_widths(table, widths_cm: list) -> None:
    """Set column widths using tblGrid XML."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for col in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(col)
    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w * 567)))  # 1cm ≈ 567 twips
        tblGrid.append(gridCol)


def build_disorder_table(doc: Document) -> None:
    """Add the disorder loop table at end of doc (will be moved later)."""
    # 2-column table: 8cm | 8cm
    table = doc.add_table(rows=10, cols=2)
    table.style = "Table Grid"
    _set_col_widths(table, [8.0, 8.0])

    # --- Row 0: loop start control row (minimal height) ---
    r0 = table.rows[0]
    _set_row_height(r0, 100)
    para = r0.cells[0].paragraphs[0]
    para.add_run("{%tr for d in disorders %}")
    r0.cells[1].paragraphs[0].add_run("")

    # --- Row 1: header bar (all cols merged, IC_BLUE bg) ---
    r1 = table.rows[1]
    cell_hdr = r1.cells[0].merge(r1.cells[1])
    _set_cell_shading(cell_hdr, IC_BLUE_HEX)
    para = cell_hdr.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(para, "{{ d.ref }}", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=11)
    _add_run(para, "  –  ", color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=11)
    _add_run(para, "{{ d.name }}", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size_pt=11)

    # --- Row 2: Localisation ---
    r2 = table.rows[2]
    p2l = r2.cells[0].paragraphs[0]
    _add_run(p2l, "Localisation\xa0:", bold=True, size_pt=9)
    p2r = r2.cells[1].paragraphs[0]
    p2r.add_run("{{ d.location }}")

    # --- Row 3: Indice d'état ---
    r3 = table.rows[3]
    p3l = r3.cells[0].paragraphs[0]
    _add_run(p3l, "Indice d’\xe9tat\xa0:", bold=True, size_pt=9)
    p3r = r3.cells[1].paragraphs[0]
    p3r.add_run("{{ d.ie }}")

    # --- Row 4: Photos 1 + 2 ---
    r4 = table.rows[4]
    r4.cells[0].paragraphs[0].add_run("{{ d.photo1 }}")
    r4.cells[1].paragraphs[0].add_run("{{ d.photo2 }}")

    # --- Row 5: Photos 3 + 4 ---
    r5 = table.rows[5]
    r5.cells[0].paragraphs[0].add_run("{{ d.photo3 }}")
    r5.cells[1].paragraphs[0].add_run("{{ d.photo4 }}")

    # --- Row 6: Description ---
    r6 = table.rows[6]
    p6l = r6.cells[0].paragraphs[0]
    _add_run(p6l, "Description\xa0:", bold=True, size_pt=9)
    r6.cells[1].paragraphs[0].add_run("{{ d.description }}")

    # --- Row 7: Cause ---
    r7 = table.rows[7]
    p7l = r7.cells[0].paragraphs[0]
    _add_run(p7l, "Cause probable\xa0:", bold=True, size_pt=9)
    r7.cells[1].paragraphs[0].add_run("{{ d.cause }}")

    # --- Row 8: Recommandations ---
    r8 = table.rows[8]
    p8l = r8.cells[0].paragraphs[0]
    _add_run(p8l, "Recommandations\xa0:", bold=True, size_pt=9)
    r8.cells[1].paragraphs[0].add_run("{{ d.recommendations }}")

    # --- Row 9: loop end control row ---
    r9 = table.rows[9]
    _set_row_height(r9, 100)
    r9.cells[0].paragraphs[0].add_run("{%tr endfor %}")
    r9.cells[1].paragraphs[0].add_run("")

    return table


def main() -> None:
    print(f"Source : {SOURCE}")
    print(f"Output : {OUTPUT}")

    doc = Document(str(SOURCE))
    body = doc.element.body
    children = list(body)

    # Find anchor paragraphs
    obs_heading_el = None
    synth_heading_el = None
    for p in doc.paragraphs:
        if "Observations et désordres" in p.text and p.style.name.startswith("Heading"):
            obs_heading_el = p._element
        if "Synthèse et recommandations" in p.text and p.style.name.startswith("Heading"):
            synth_heading_el = p._element

    if obs_heading_el is None:
        raise SystemExit("ERROR: Could not find 'Observations et désordres' heading in source template.")
    if synth_heading_el is None:
        raise SystemExit("ERROR: Could not find 'Synthèse et recommandations' heading in source template.")

    obs_idx = children.index(obs_heading_el)
    synth_idx = children.index(synth_heading_el)
    print(f"Removing body elements [{obs_idx}:{synth_idx}] (observations section)...")

    # Remove all elements from obs_heading to (but not including) synth_heading
    to_remove = children[obs_idx:synth_idx]
    for el in to_remove:
        body.remove(el)

    # Rebuild children list after removal
    children = list(body)
    synth_idx = children.index(synth_heading_el)

    # Add new "Observations et désordres" heading before synthesis
    new_heading = doc.add_heading("Observations et désordres", level=1)
    new_heading_el = new_heading._element
    # Move from end to before synth heading
    body.remove(new_heading_el)
    body.insert(synth_idx, new_heading_el)
    synth_idx += 1

    # Build the disorder loop table (added at end temporarily)
    table = build_disorder_table(doc)
    tbl_el = table._tbl
    # Move from end to before synth heading
    body.remove(tbl_el)
    body.insert(synth_idx, tbl_el)
    synth_idx += 1

    # Add a blank paragraph after the table (before synthesis)
    blank_para = OxmlElement("w:p")
    body.insert(synth_idx, blank_para)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"✅ Saved: {OUTPUT}")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
