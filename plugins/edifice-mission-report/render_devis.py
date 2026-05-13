#!/usr/bin/env python3
"""
Render a Rapport Préliminaire / Devis DOCX for IC Ingénieurs Conseils.

This renderer produces the internal proposal/estimate document used before
a mission is formally engaged. Structure follows the standard IC intake form:
1 — Contexte client
2 — Mission
3 — Bâtiment
4 — Documents fournis
5 — Observations terrain
6 — Proposition de mission
7 — Chiffrage estimatif
Validation

Usage:
    python render_devis.py context.json [--output devis.docx]

Context JSON schema:
{
    "project_type": "devis",
    "titre_service": "Diagnostic structurel planchers bois",
    "client": "Nom du client",
    "type_acteur": "Syndic",            # MOA / MOE / Entrepreneur / Syndic / Particulier
    "interlocuteur_nom": "Jean Dupont",
    "interlocuteur_role": "Gestionnaire de copropriété",
    "interlocuteur_contact": "jean@example.com — 06 XX XX XX XX",
    "adresse_client": "15 rue des Lilas, 75010 Paris",
    "type_mission": "Diagnostic",       # Diagnostic / Exécution / CCTP / Calcul / AMO
    "declencheur": "Lors de travaux de rénovation, des poutres dégradées ont été découvertes.",
    "livrable": "Rapport DOCX avec observations et recommandations",
    "urgence": "Normal",                # Normal / Court terme / Urgent
    "adresse": "46 Rue de Varenne, 75007 Paris",
    "type_batiment": "Immeuble résidentiel",
    "annee_construction": "XVIIIe siècle",
    "nb_etages": "R+2 + combles",
    "surface": "",
    "description_batiment": "Hôtel particulier...",
    "documents_fournis": [
        {"document": "Plans architecte / plans existants", "fourni": false},
        {"document": "DOE", "fourni": false}
    ],
    "observations": [
        {
            "localisation": "Chambre 1",
            "desordre": "Linteau dégradé, pourriture avancée",
            "donnees_cles": "L=2,20m — 25×25cm",
            "ref_photo": "P1"
        }
    ],
    "proposition_mission": "Diagnostic structurel complet du plancher bois...",
    "incertitudes": "Accès aux zones masquées conditionné à la dépose des faux-plafonds.",
    "chiffrage": [
        {"prestation": "Déplacement terrain", "nb_heures": "2", "montant_ht": ""},
        {"prestation": "Étude des documents fournis", "nb_heures": "1", "montant_ht": ""},
        {"prestation": "Visite terrain", "nb_heures": "3", "montant_ht": ""},
        {"prestation": "Rédaction du rapport", "nb_heures": "4", "montant_ht": ""}
    ],
    "technicien": "R. Laborbe",
    "date_visite": "2026-04-28",
    "date_envoi": ""
}
"""

import argparse
import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

# IC Ingénieurs brand colors
IC_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
IC_BLUE_LIGHT = RGBColor(0xE8, 0xEF, 0xF8)
IC_ORANGE = RGBColor(0xE8, 0x6A, 0x1A)
IC_HEADER_BG = "1F3A5F"
IC_ROW_ALT = "F0F4F8"

MOIS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]

# Default documents checklist when none provided
DEFAULT_DOCUMENTS = [
    "Plans architecte / plans existants",
    "DGT (Documents Graphiques Techniques)",
    "DOE (Dossier des Ouvrages Exécutés)",
    "Rapports d'études antérieurs",
    "Permis de construire",
    "Photos existantes",
]

# Default chiffrage lines
DEFAULT_CHIFFRAGE = [
    "Déplacement(s) terrain",
    "Étude des documents fournis",
    "Visite terrain",
    "Rédaction du rapport / note de calcul",
    "Réunion / coordination",
]


def format_date_french(iso_date: str) -> str:
    try:
        dt = datetime.strptime(iso_date[:10], "%Y-%m-%d")
        return f"{dt.day} {MOIS_FR[dt.month - 1]} {dt.year}"
    except (ValueError, TypeError):
        return iso_date or ""


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_table_borders(table, color: str = "AAAAAA") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_row_2col(table, label: str, value: str, shade_label: bool = True) -> None:
    """Add a label / value row to a 2-column table."""
    row = table.add_row()
    lbl_cell, val_cell = row.cells[0], row.cells[1]

    lbl_cell.text = label
    if lbl_cell.paragraphs[0].runs:
        r = lbl_cell.paragraphs[0].runs[0]
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = IC_BLUE
    if shade_label:
        set_cell_shading(lbl_cell, "EEF2F8")
    set_cell_margins(lbl_cell, 50, 50, 100, 80)

    val_cell.text = value
    if val_cell.paragraphs[0].runs:
        val_cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_margins(val_cell, 50, 50, 100, 80)


def add_section_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(f"{number} — ")
    run_num.font.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = IC_ORANGE
    run_title = p.add_run(title.upper())
    run_title.font.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = IC_BLUE


# ── Cover ─────────────────────────────────────────────────────────────────────

def build_cover(doc: Document, ctx: dict) -> None:
    doc.add_paragraph()

    # IC Ingénieurs header line
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_p.add_run("IC INGÉNIEURS CONSEILS")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = IC_BLUE
    r.font.all_caps = True

    doc.add_paragraph()

    # Document type badge
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_run = badge_p.add_run("RAPPORT PRÉLIMINAIRE — DEVIS")
    badge_run.font.size = Pt(9)
    badge_run.font.bold = True
    badge_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Simulate background with shading via a 1-cell table
    badge_table = doc.add_table(rows=1, cols=1)
    set_cell_shading(badge_table.rows[0].cells[0], IC_HEADER_BG)
    badge_table.rows[0].cells[0].text = ""
    p_badge = badge_table.rows[0].cells[0].paragraphs[0]
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_badge.add_run("RAPPORT PRÉLIMINAIRE — DEMANDE DE DEVIS")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_margins(badge_table.rows[0].cells[0], 120, 120, 200, 200)

    doc.add_paragraph()

    # Mission title
    titre = ctx.get("titre_service") or "Diagnostic / Étude de structure"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(titre)
    tr.font.size = Pt(16)
    tr.font.bold = True
    tr.font.color.rgb = IC_BLUE

    doc.add_paragraph()

    # Key info block (client, adresse, date)
    meta = [
        ("Client", ctx.get("client") or ""),
        ("Adresse", ctx.get("adresse") or ""),
        ("Date visite", format_date_french(ctx.get("date_visite") or "")),
        ("Technicien", ctx.get("technicien") or ""),
    ]
    for lbl, val in meta:
        if not val:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{lbl} : ").font.bold = True
        p.runs[-1].font.size = Pt(10)
        p.runs[-1].font.color.rgb = IC_BLUE
        r_val = p.add_run(val)
        r_val.font.size = Pt(10)

    doc.add_page_break()


# ── Section 1 : Contexte client ───────────────────────────────────────────────

def build_section_client(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "1", "Contexte client")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = "Champ"
    hdr.cells[1].text = "Information"
    for cell in hdr.cells:
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    rows = [
        ("Client / donneur d'ordre", ctx.get("client") or ""),
        ("Type d'acteur", ctx.get("type_acteur") or ""),
        ("Interlocuteur — Nom", ctx.get("interlocuteur_nom") or ""),
        ("Interlocuteur — Rôle", ctx.get("interlocuteur_role") or ""),
        ("Contact (email / tél.)", ctx.get("interlocuteur_contact") or ""),
        ("Adresse client", ctx.get("adresse_client") or ""),
    ]
    for label, value in rows:
        add_row_2col(table, label, value)

    doc.add_paragraph()


# ── Section 2 : Mission ───────────────────────────────────────────────────────

def build_section_mission(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "2", "Mission")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0]
    hdr.cells[0].text = "Champ"
    hdr.cells[1].text = "Information"
    for cell in hdr.cells:
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    rows = [
        ("Type de mission", ctx.get("type_mission") or ""),
        ("Déclencheur (contexte)", ctx.get("declencheur") or ""),
        ("Livrable attendu", ctx.get("livrable") or ""),
        ("Urgence", ctx.get("urgence") or "Normal"),
    ]
    for label, value in rows:
        add_row_2col(table, label, value)

    doc.add_paragraph()


# ── Section 3 : Bâtiment ──────────────────────────────────────────────────────

def build_section_batiment(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "3", "Bâtiment")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0]
    hdr.cells[0].text = "Caractéristique"
    hdr.cells[1].text = "Information"
    for cell in hdr.cells:
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    rows = [
        ("Adresse complète", ctx.get("adresse") or ""),
        ("Type de bâtiment", ctx.get("type_batiment") or ""),
        ("Année de construction", ctx.get("annee_construction") or ""),
        ("Nombre d'étages", ctx.get("nb_etages") or ""),
        ("Surface approximative", ctx.get("surface") or ""),
        ("Description", ctx.get("description_batiment") or ""),
    ]
    for label, value in rows:
        add_row_2col(table, label, value)

    doc.add_paragraph()


# ── Section 4 : Documents fournis ────────────────────────────────────────────

def build_section_documents(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "4", "Documents fournis par le client")

    docs_list = ctx.get("documents_fournis")
    if not docs_list:
        # Build default list with all "Non"
        docs_list = [{"document": d, "fourni": False} for d in DEFAULT_DOCUMENTS]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0]
    hdr.cells[0].text = "Document"
    hdr.cells[1].text = "Fourni ?"
    for cell in hdr.cells:
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    for i, item in enumerate(docs_list):
        row = table.add_row()
        doc_name = item.get("document") or item if isinstance(item, str) else ""
        fourni = item.get("fourni") if isinstance(item, dict) else False
        fourni_str = "✓ Oui" if fourni else "✗ Non"

        row.cells[0].text = doc_name
        row.cells[1].text = fourni_str
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        if row.cells[1].paragraphs[0].runs:
            r = row.cells[1].paragraphs[0].runs[0]
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x00, 0x80, 0x00) if fourni else RGBColor(0xAA, 0xAA, 0xAA)
        for cell in row.cells:
            set_cell_margins(cell, 40, 40, 100, 80)
            if i % 2 == 0:
                set_cell_shading(cell, IC_ROW_ALT)

    doc.add_paragraph()


# ── Section 5 : Observations terrain ─────────────────────────────────────────

def build_section_observations(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "5", "Observations terrain")

    note_p = doc.add_paragraph()
    note_run = note_p.add_run(
        "ℹ  Les photos sont insérées en annexe numérotées P1, P2… et référencées dans la colonne « Réf. photo(s) » ci-dessous."
    )
    note_run.font.size = Pt(8)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    observations = ctx.get("observations") or []

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_borders(table)

    headers = ["Localisation", "Désordre / Observation", "Données clés (mesures, section…)", "Réf. photo(s)"]
    widths_dxa = [2500, 5000, 3000, 1200]
    hdr = table.rows[0]
    for cell, h, w in zip(hdr.cells, headers, widths_dxa):
        cell.text = h
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(8)
        set_cell_margins(cell, 40, 40, 60, 60)
        # Set column width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{w}" w:type="dxa"/>')
        tcPr.append(tcW)

    # Ensure at least 5 empty rows for fill-in if no observations
    obs_to_render = observations if observations else [{}] * 5

    for i, obs in enumerate(obs_to_render):
        row = table.add_row()
        cells = row.cells

        values = [
            obs.get("localisation") or "",
            obs.get("desordre") or "",
            obs.get("donnees_cles") or "",
            obs.get("ref_photo") or "",
        ]
        for cell, val in zip(cells, values):
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_margins(cell, 50, 50, 60, 60)
            if i % 2 == 0:
                set_cell_shading(cell, IC_ROW_ALT)

    doc.add_paragraph()


# ── Section 6 : Proposition de mission ───────────────────────────────────────

def build_section_proposition(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "6", "Proposition de mission")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0]
    hdr.cells[0].text = "Champ"
    hdr.cells[1].text = "Contenu"
    for cell in hdr.cells:
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    add_row_2col(table, "Prestations proposées", ctx.get("proposition_mission") or "")
    add_row_2col(table, "Incertitudes / réserves", ctx.get("incertitudes") or "")

    doc.add_paragraph()


# ── Section 7 : Chiffrage estimatif ──────────────────────────────────────────

def build_section_chiffrage(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "7", "Chiffrage estimatif")

    chiffrage = ctx.get("chiffrage")
    if not chiffrage:
        chiffrage = [{"prestation": p, "nb_heures": "", "montant_ht": ""} for p in DEFAULT_CHIFFRAGE]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)

    hdr = table.rows[0]
    headers = ["Prestation", "Nb heures est.", "Montant HT est."]
    for cell, h in zip(hdr.cells, headers):
        cell.text = h
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    total_hours = 0
    for i, item in enumerate(chiffrage):
        row = table.add_row()
        cells = row.cells

        prestation = item.get("prestation") or ""
        nb_h = str(item.get("nb_heures") or "")
        montant = str(item.get("montant_ht") or "")

        cells[0].text = prestation
        cells[1].text = nb_h
        cells[2].text = montant

        for j, cell in enumerate(cells):
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(9)
                if j > 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(cell, 40, 40, 80, 80)
            if i % 2 == 0:
                set_cell_shading(cell, IC_ROW_ALT)

        try:
            total_hours += float(nb_h)
        except (ValueError, TypeError):
            pass

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = str(int(total_hours)) if total_hours else ""
    total_row.cells[2].text = ""
    for cell in total_row.cells:
        set_cell_shading(cell, "E8EFF8")
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = IC_BLUE
        set_cell_margins(cell, 40, 40, 80, 80)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if cell != total_row.cells[0] else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


# ── Validation ────────────────────────────────────────────────────────────────

def build_section_validation(doc: Document, ctx: dict) -> None:
    add_section_heading(doc, "✓", "Validation")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)

    headers = ["Technicien", "Date visite terrain", "Date envoi à Laurent"]
    values = [
        ctx.get("technicien") or "",
        format_date_french(ctx.get("date_visite") or ""),
        format_date_french(ctx.get("date_envoi") or "") if ctx.get("date_envoi") else "",
    ]

    hdr_row = table.rows[0]
    for cell, h in zip(hdr_row.cells, headers):
        cell.text = h
        set_cell_shading(cell, IC_HEADER_BG)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        set_cell_margins(cell, 40, 40, 100, 80)

    val_row = table.rows[1]
    for cell, val in zip(val_row.cells, values):
        cell.text = val
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_margins(cell, 60, 60, 100, 80)

    doc.add_paragraph()


# ── Main renderer ─────────────────────────────────────────────────────────────

def render_devis(context: dict, photos_dir: str = ".", output_path: str = "rapport_devis.docx") -> str:
    """Render a Rapport Préliminaire / Devis DOCX from context data."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    build_cover(doc, context)
    build_section_client(doc, context)
    build_section_mission(doc, context)
    build_section_batiment(doc, context)
    build_section_documents(doc, context)
    build_section_observations(doc, context)
    build_section_proposition(doc, context)
    build_section_chiffrage(doc, context)
    build_section_validation(doc, context)

    doc.save(output_path)
    obs_count = len(context.get("observations") or [])
    print(f"✅ Rapport devis saved: {output_path}")
    print(f"   {obs_count} observation(s)")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Render Rapport Préliminaire / Devis DOCX")
    parser.add_argument("context", help="Path to context JSON file")
    parser.add_argument("--output", default="rapport_devis.docx", help="Output path")
    args = parser.parse_args()

    with open(args.context, encoding="utf-8") as f:
        ctx = json.load(f)

    render_devis(ctx, output_path=args.output)
