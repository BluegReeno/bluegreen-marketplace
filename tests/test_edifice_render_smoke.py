"""
Deterministic smoke test for the edifice report pipeline: build_context.py -> render_report.py.

Covers both project types (diagnostic, suivi_chantier) against fixed fixtures under
tests/fixtures/edifice/<project_type>/ — mirrors what get_mission_with_assets (mcp_response.json)
and the improve step (context.json) produce. No Supabase, no network, no LLM call.

Assertions are structural invariants (observation count, image count delta, key headers),
never exact .docx bytes — docxtpl re-encodes images non-deterministically on every run.

Closes #44.
"""
import json
import pathlib
import subprocess
import sys
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

REPO_ROOT = pathlib.Path(__file__).parent.parent
SCRIPTS_DIR = REPO_ROOT / "plugins/hal/scripts"
TEMPLATES_DIR = REPO_ROOT / "plugins/hal/templates/ic-ingenieurs"
FIXTURES_DIR = pathlib.Path(__file__).parent / "fixtures/edifice"

sys.path.insert(0, str(SCRIPTS_DIR))
import build_context  # noqa: E402


def _run_build_context(project_type: str, output_dir: pathlib.Path) -> dict:
    mcp_response = FIXTURES_DIR / project_type / "mcp_response.json"
    argv = ["build_context.py", str(mcp_response), str(output_dir)]
    with patch.object(sys, "argv", argv):
        build_context.main()
    return json.loads((output_dir / "context.json").read_text(encoding="utf-8"))


def _full_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _render(project_type: str, context_dir: pathlib.Path, output_path: pathlib.Path) -> None:
    result = subprocess.run(
        [
            sys.executable, str(SCRIPTS_DIR / "render_report.py"),
            str(context_dir / "context.json"),
            "--photos-dir", str(context_dir / "photos"),
            "--output", str(output_path),
        ],
        capture_output=True, text=True,
    )
    assert result.returncode == 0, f"render_report.py failed for {project_type}:\n{result.stderr}"
    assert output_path.exists(), f"render_report.py did not write {output_path}"


class TestBuildContextDiagnostic(unittest.TestCase):
    """build_context.py on a fixed mcp_response.json -> context.json shape (diagnostic)."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.ctx = _run_build_context("diagnostic", pathlib.Path(self._tmp.name) / "mission")

    def tearDown(self):
        self._tmp.cleanup()

    def test_notes_routed_to_observations_by_type(self):
        # 2 "disorder" notes -> observations[], 1 "note" (methodo-tagged) -> notes[]
        self.assertEqual(len(self.ctx["observations"]), 2)
        self.assertEqual(len(self.ctx["notes"]), 1)

    def test_assessment_preserved_on_diagnostic_1_4_scale(self):
        by_ref = {o["ref"]: o for o in self.ctx["observations"]}
        self.assertEqual(by_ref["OBS-01"]["assessment"], "2")
        self.assertEqual(by_ref["OBS-02"]["assessment"], "1")

    def test_photo_filenames_resolved_from_storage_path(self):
        by_ref = {o["ref"]: o for o in self.ctx["observations"]}
        self.assertEqual(by_ref["OBS-01"]["photos"], ["diag_photo_1.jpg"])
        self.assertEqual(by_ref["OBS-01"]["photo"], "diag_photo_1.jpg")
        self.assertEqual(by_ref["OBS-02"]["photos"], ["diag_photo_2.jpg"])

    def test_methodo_tag_preserved_on_free_note(self):
        self.assertEqual(self.ctx["notes"][0]["metadata"]["tag"], "methodo:visite_terrain")
        self.assertEqual(self.ctx["notes"][0]["photo"], "methodo_1.jpg")

    def test_header_fields_mapped_and_address_cleaned(self):
        self.assertEqual(self.ctx["residence"], "Résidence Les Tilleuls")
        # _clean_address strips the BAN geocoder parenthetical suffix
        self.assertEqual(self.ctx["adresse"], "12 Rue de Test 75012")
        self.assertEqual(self.ctx["ref_dossier"], "DIAG-2026-001")
        self.assertEqual(self.ctx["date_visite"], "2026-03-15")


class TestBuildContextSuiviChantier(unittest.TestCase):
    """build_context.py on a fixed mcp_response.json -> context.json shape (suivi_chantier)."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.ctx = _run_build_context("suivi_chantier", pathlib.Path(self._tmp.name) / "mission")

    def tearDown(self):
        self._tmp.cleanup()

    def test_reservation_notes_routed_to_observations(self):
        self.assertEqual(len(self.ctx["observations"]), 2)
        self.assertEqual(len(self.ctx["notes"]), 0)

    def test_assessment_preserved_on_suivi_chantier_states(self):
        by_ref = {o["ref"]: o for o in self.ctx["observations"]}
        self.assertEqual(by_ref["V-01"]["assessment"], "a_faire")
        self.assertEqual(by_ref["V-02"]["assessment"], "observation")

    def test_photo_filenames_resolved_from_storage_path(self):
        by_ref = {o["ref"]: o for o in self.ctx["observations"]}
        self.assertEqual(by_ref["V-01"]["photo"], "cr_photo_1.jpg")
        self.assertEqual(by_ref["V-02"]["photo"], "cr_photo_2.jpg")

    def test_header_fields_mapped_and_address_cleaned(self):
        self.assertEqual(self.ctx["residence"], "Résidence Le Gros Saule")
        self.assertEqual(self.ctx["adresse"], "5 Avenue du Chantier 93600")
        self.assertEqual(self.ctx["ref_dossier"], "CR-2026-014")


class TestRenderDiagnosticSmoke(unittest.TestCase):
    """render_report.py on a fixed, already-improved context.json -> invariants on the .docx."""

    def test_docx_opens_with_expected_disorders_images_and_headers(self):
        context_dir = FIXTURES_DIR / "diagnostic"
        with tempfile.TemporaryDirectory() as tmp:
            output_path = pathlib.Path(tmp) / "rapport.docx"
            _render("diagnostic", context_dir, output_path)

            doc = Document(str(output_path))
            text = _full_text(doc)

            # 2 disorders -> each disorder renders as its own table, keyed off a stable
            # template label ("Localisation :"), independent of unrelated table changes.
            localisation_tables = [
                t for t in doc.tables
                if t.rows and t.rows[0].cells[0].text.strip() == "Localisation\xa0:"
            ]
            self.assertEqual(len(localisation_tables), 2)
            self.assertIn("OBS-01", text)
            self.assertIn("OBS-02", text)
            self.assertIn("Fissure façade nord", text)
            self.assertIn("Affaissement plancher", text)

            # 3 fixture photos (2 disorders x1 + 1 methodo) on top of the template's own
            # static artwork (logos/letterhead) — compare against the pristine template.
            baseline_shapes = len(Document(str(TEMPLATES_DIR / "diagnostic.docx")).inline_shapes)
            self.assertEqual(len(doc.inline_shapes) - baseline_shapes, 3)

            # Key headers
            self.assertIn("Résidence Les Tilleuls", text)
            self.assertIn("12 Rue de Test", text)
            self.assertIn("15 mars 2026", text)


class TestRenderSuiviChantierSmoke(unittest.TestCase):
    """render_report.py on a fixed, already-improved context.json -> invariants on the .docx."""

    def test_docx_opens_with_expected_observations_images_and_headers(self):
        context_dir = FIXTURES_DIR / "suivi_chantier"
        with tempfile.TemporaryDirectory() as tmp:
            output_path = pathlib.Path(tmp) / "cr.docx"
            _render("suivi_chantier", context_dir, output_path)

            doc = Document(str(output_path))
            text = _full_text(doc)

            observations_table = next(
                t for t in doc.tables
                if t.rows and t.rows[0].cells[0].text.strip() == "Étage / Façade"
            )
            self.assertEqual(len(observations_table.rows) - 1, 2)  # header row + 2 observations
            self.assertIn("Recouvrement des aciers", text)
            self.assertIn("Espacement des cadres", text)

            # 2 fixture photos (1 per observation); template has no static artwork.
            baseline_shapes = len(Document(str(TEMPLATES_DIR / "suivi_chantier.docx")).inline_shapes)
            self.assertEqual(len(doc.inline_shapes) - baseline_shapes, 2)

            # Key headers
            self.assertIn("Résidence Le Gros Saule", text)
            self.assertIn("5 Avenue du Chantier", text)
            self.assertIn("17 avril 2026", text)


if __name__ == "__main__":
    unittest.main()
